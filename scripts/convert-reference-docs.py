#!/usr/bin/env python3
"""
Regenerate docs/reference/*.md from the authored .docx originals.

CLAUDE.md points build sessions at these specs by Markdown filename, and
Markdown is what greps and diffs in review — but the .docx files are what the
founder actually edits. This keeps the two in step.

The Markdown is generated output: do not hand-edit it, because the next run
overwrites it. Corrections belong in the .docx. Where the build knowingly
diverges from a spec, that belongs in an ADR, not an edit to either.

Usage:
    python3 scripts/convert-reference-docs.py

Requires python-docx (`pip install python-docx`).
"""

import pathlib
import re
import sys

try:
    import docx
except ImportError:
    sys.exit("python-docx is not installed — run: pip install python-docx")

REFERENCE_DIR = pathlib.Path(__file__).resolve().parent.parent / "docs" / "reference"

# Source .docx -> generated .md. The Markdown names are the ones CLAUDE.md
# references, so renaming one here means updating CLAUDE.md too.
DOCUMENTS = {
    "VirtualOfficeAI_01_VersioningLedger.docx": "Versioning_Ledger.md",
    "VirtualOfficeAI_02_StrategicVision_v4.docx": "Strategic_Vision.md",
    "VirtualOfficeAI_03_PlatformSpecification_v2.docx": "Platform_Specification.md",
    "VirtualOfficeAI_04_SystemArchitecture_v2.docx": "System_Architecture.md",
    "VirtualOfficeAI_05_UXSpecification_v2.docx": "UX_Specification.md",
    "VirtualOfficeAI_06_AdminConsoleSpecification.docx": "Admin_Console_Specification.md",
    "VirtualOfficeAI_07_SprintPlan.docx": "Sprint_Plan.md",
    "VirtualOfficeAI_08_SessionOperatingManual.docx": "Session_Operating_Manual.md",
    "VirtualOfficeAI_09_GTMStrategy.docx": "GTM_Strategy.md",
}

BANNER = (
    "<!-- Generated from {src} by scripts/convert-reference-docs.py.\n"
    "     Do not hand-edit: edit the .docx and regenerate. -->\n"
)


def style_name(paragraph) -> str:
    """Paragraphs with no explicit style have `style` set to None."""
    try:
        return (paragraph.style.name or "").lower()
    except AttributeError:
        return ""


def cell(text: str) -> str:
    """Table cells may contain newlines and pipes, both of which break a row."""
    return text.strip().replace("\n", " ").replace("|", "\\|")


def convert(path: pathlib.Path) -> str:
    document = docx.Document(str(path))
    blocks: list[str] = []
    paragraphs, tables = document.paragraphs, document.tables
    para_i = tbl_i = 0

    # Walk the body in document order so tables stay positioned relative to the
    # prose around them; python-docx's .paragraphs and .tables are separate
    # sequences and reading them in turn would move every table to the end.
    for child in document.element.body.iterchildren():
        tag = child.tag.split("}")[-1]

        if tag == "p" and para_i < len(paragraphs):
            paragraph = paragraphs[para_i]
            para_i += 1
            text = paragraph.text.strip()
            if not text:
                continue
            style = style_name(paragraph)
            heading = re.match(r"heading (\d)", style)
            if heading:
                blocks.append("#" * min(int(heading.group(1)), 6) + " " + text)
            elif style.startswith("list"):
                blocks.append("- " + text)
            else:
                blocks.append(text)

        elif tag == "tbl" and tbl_i < len(tables):
            table = tables[tbl_i]
            tbl_i += 1
            rows = [[cell(c.text) for c in r.cells] for r in table.rows]
            if not rows:
                continue
            width = max(len(r) for r in rows)
            rows = [r + [""] * (width - len(r)) for r in rows]
            blocks.append("| " + " | ".join(rows[0]) + " |")
            blocks.append("|" + "---|" * width)
            blocks.extend("| " + " | ".join(r) + " |" for r in rows[1:])

    return "\n\n".join(blocks) + "\n"


def main() -> int:
    missing = [s for s in DOCUMENTS if not (REFERENCE_DIR / s).exists()]
    if missing:
        print("missing source documents:", ", ".join(missing), file=sys.stderr)

    for source, target in DOCUMENTS.items():
        src = REFERENCE_DIR / source
        if not src.exists():
            continue
        markdown = BANNER.format(src=source) + "\n" + convert(src)
        (REFERENCE_DIR / target).write_text(markdown)
        headings = len(re.findall(r"^#", markdown, re.M))
        print(f"{target}: {len(markdown.splitlines())} lines, {headings} headings")

    return 1 if missing else 0


if __name__ == "__main__":
    raise SystemExit(main())
